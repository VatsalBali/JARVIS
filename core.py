"""
ORACLE - Day 1: Talk to an LLM (via Groq's free API) with tool calling.

Why Groq and not a local model: local inference needs a real GPU. With
integrated graphics only, running even a "small" model on CPU is painfully
slow. Groq gives free, fast, hosted inference instead - and it's
OpenAI-compatible, so the code below looks almost identical to what you'd
write for OpenAI or any other compatible provider.

SETUP:
1. Get a free API key: https://console.groq.com  (no credit card needed)
2. Set it as an environment variable so it's never hardcoded in this file:
       export GROQ_API_KEY="your-key-here"      (Mac/Linux)
       setx GROQ_API_KEY "your-key-here"         (Windows, new terminal after)
3. Install deps:      pip install groq
4. Run this file:     python core.py
"""

from groq import Groq
from tavily import TavilyClient
import json
import os
import sys
import sqlite3
import subprocess
import shutil
import webbrowser
from urllib.parse import quote_plus
import msal
import requests
import base64
from email.mime.text import MIMEText
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request as GoogleAuthRequest
from google.oauth2.credentials import Credentials
import psutil
import sounddevice as sd
import numpy as np
import threading
import re
from pathlib import Path
from datetime import datetime, timedelta

MODEL = "openai/gpt-oss-120b"
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

SYSTEM_PROMPT = (
    "You are ORACLE, a personal AI assistant in the spirit of JARVIS from Iron Man: "
    "calm, dry-witted, quietly loyal, and understated rather than effusive. A touch of "
    "wit is welcome where it fits naturally, but never at the expense of being direct "
    "and efficient when the user actually needs something done - personality seasons "
    "your responses, it doesn't pad them out. When a question needs real information "
    "about this machine or the world, call a tool instead of guessing. To open an "
    "application (like Notepad, Chrome, or Spotify), call launch_app directly with the "
    "app name - do not use list_files or open_file to search for it first. For "
    "non-trivial coding help - writing code, debugging, explaining code, "
    "architecture questions - use ask_coding_agent to consult a coding "
    "specialist rather than answering directly yourself."
)

# Name ORACLE uses when greeting you on the ring-click voice flow (or generate_wake_greeting, if reused later).
USER_NAME = "Oracle"

# A directory listing goes into the conversation and is re-sent on every later
# turn, so an uncapped one (System32 is ~23k tokens) blows the API rate limit.
MAX_LISTED_FILES = 50

MAX_TOOL_ROUNDS = 5

def _get_data_dir() -> str:
    """
    Returns a proper per-user, writable directory for ORACLE.s persistent
    data (currently just the SQLite DB). A relative path like
    "oracle_memory.db" only works reliably when you run `python core.py`
    by hand from this exact folder - it breaks for a packaged .exe
    launched via Windows autostart, which often runs with an unexpected
    working directory (frequently System32). %LOCALAPPDATA%\\ORACLE is
    the standard place per-user app data belongs on Windows.
    """
    if sys.platform == "win32":
        base = os.environ.get("LOCALAPPDATA") or os.path.expanduser("~")
    else:
        base = os.path.expanduser("~/.local/share")
    data_dir = os.path.join(base, "ORACLE")
    os.makedirs(data_dir, exist_ok=True)
    return data_dir


DB_PATH = os.path.join(_get_data_dir(), "oracle_memory.db")

# How many of the most recent messages we actually SEND to the model each
# turn. Everything is still saved to disk - this only limits what gets
# resent on each API call, which is what keeps token usage (and cost/rate
# limits) from growing unbounded as a conversation gets long.
MAX_HISTORY_MESSAGES = 20


# ---------------------------------------------------------------------------
# PERSISTENCE: save every message to SQLite so ORACLE remembers past
# conversations even after the script exits and restarts. This is the
# "item #5 - conversation memory" piece leveling up from a plain in-RAM list
# to something durable.
# ---------------------------------------------------------------------------

def init_db():
    """Creates all tables if they don't already exist, and migrates older
    databases (from before conversation threading, or before projects)
    by adding new columns if missing."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS conversations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER,
            role TEXT NOT NULL,
            content TEXT,
            tool_calls TEXT,
            tool_call_id TEXT,
            timestamp TEXT NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS projects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            path TEXT NOT NULL UNIQUE,
            name TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT
        )
    """)
    # Migration: a DB created before conversation threading existed won't
    # have this column - old messages just end up with conversation_id
    # NULL (not deleted, just invisible to the new per-thread sidebar).
    cols = [row[1] for row in conn.execute("PRAGMA table_info(messages)").fetchall()]
    if "conversation_id" not in cols:
        conn.execute("ALTER TABLE messages ADD COLUMN conversation_id INTEGER")
    # Migration: a DB created before projects existed won't have this
    # column on conversations - old conversations just end up with
    # project_id NULL (regular chats, not tied to any project).
    conv_cols = [row[1] for row in conn.execute("PRAGMA table_info(conversations)").fetchall()]
    if "project_id" not in conv_cols:
        conn.execute("ALTER TABLE conversations ADD COLUMN project_id INTEGER")
    if "pinned" not in conv_cols:
        conn.execute("ALTER TABLE conversations ADD COLUMN pinned INTEGER DEFAULT 0")

    proj_cols = [row[1] for row in conn.execute("PRAGMA table_info(projects)").fetchall()]
    if "pinned" not in proj_cols:
        conn.execute("ALTER TABLE projects ADD COLUMN pinned INTEGER DEFAULT 0")

    conn.commit()
    conn.close()


def create_conversation(first_message: str, project_id: int = None) -> int:
    """
    Starts a new conversation thread, titled from a truncated version of
    its first message (simple and free - no extra LLM call just to name
    a chat). If project_id is given, this thread is tied to that project
    rather than being a regular standalone chat. Returns the new
    conversation's ID.
    """
    title = first_message.strip().replace("\n", " ")[:50]
    if len(first_message.strip()) > 50:
        title += "..."
    now = datetime.now().isoformat()
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.execute(
        "INSERT INTO conversations (title, created_at, updated_at, project_id) VALUES (?, ?, ?, ?)",
        (title, now, now, project_id),
    )
    conn.commit()
    conversation_id = cursor.lastrowid
    conn.close()
    return conversation_id


def get_setting(key: str) -> str:
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT value FROM settings WHERE key = ?", (key,)).fetchone()
    conn.close()
    return row[0] if row else None


def set_setting(key: str, value: str):
    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        "INSERT INTO settings (key, value) VALUES (?, ?) "
        "ON CONFLICT(key) DO UPDATE SET value = excluded.value",
        (key, value),
    )
    conn.commit()
    conn.close()


def get_or_create_project(path: str) -> int:
    """Looks up a project by its folder path, creating one if this is the
    first time it's been opened. Returns the project's ID."""
    name = os.path.basename(os.path.normpath(path)) or path
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT id FROM projects WHERE path = ?", (path,)).fetchone()
    if row:
        conn.close()
        return row[0]
    now = datetime.now().isoformat()
    cursor = conn.execute(
        "INSERT INTO projects (path, name, created_at) VALUES (?, ?, ?)",
        (path, name, now),
    )
    conn.commit()
    project_id = cursor.lastrowid
    conn.close()
    return project_id


def get_project_conversation_id(project_id: int):
    """Returns the existing conversation thread for a project, if one has
    been started yet - each project has exactly one ongoing thread, not
    multiple named chats like the regular sidebar."""
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT id FROM conversations WHERE project_id = ? ORDER BY updated_at DESC LIMIT 1",
        (project_id,),
    ).fetchone()
    conn.close()
    return row[0] if row else None


def get_project(project_id: int) -> dict:
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT id, path, name, pinned FROM projects WHERE id = ?", (project_id,)
    ).fetchone()
    conn.close()
    if not row:
        return None
    return {"id": row[0], "path": row[1], "name": row[2], "pinned": bool(row[3])}


def list_projects_in_root(root_path: str) -> list:
    """
    Lists immediate subfolders of the configured projects root - always a
    live filesystem scan (not cached, so it stays accurate if folders are
    added/removed outside ORACLE), merged with each folder's tracked
    record (a possibly-renamed display name, pinned state) so the sidebar
    reflects both reality on disk and any customization done in ORACLE.
    Pinned projects sort first, then alphabetically by display name.
    """
    if not root_path or not os.path.isdir(root_path):
        return []
    entries = []
    try:
        for name in os.listdir(root_path):
            full_path = os.path.join(root_path, name)
            if os.path.isdir(full_path):
                project_id = get_or_create_project(full_path)
                entries.append(get_project(project_id))
    except Exception:
        pass
    entries.sort(key=lambda p: (not p["pinned"], p["name"].lower()))
    return entries


def list_conversations(limit: int = 50) -> list:
    """Returns recent STANDALONE conversations (not tied to a project),
    pinned ones first then most recently active - what populates the
    sidebar's Chats list. Project conversations show up under Projects
    instead, not duplicated here."""
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT id, title, updated_at, pinned FROM conversations "
        "WHERE project_id IS NULL ORDER BY pinned DESC, updated_at DESC LIMIT ?",
        (limit,),
    ).fetchall()
    conn.close()
    return [{"id": r[0], "title": r[1], "updated_at": r[2], "pinned": bool(r[3])} for r in rows]


def rename_conversation(conversation_id: int, new_title: str):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE conversations SET title = ? WHERE id = ?", (new_title, conversation_id))
    conn.commit()
    conn.close()


def delete_conversation(conversation_id: int):
    """Deletes a chat and its messages from ORACLE's memory. This only
    ever touches ORACLE's own database - never any files on your
    computer, even for project-linked conversations."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("DELETE FROM messages WHERE conversation_id = ?", (conversation_id,))
    conn.execute("DELETE FROM conversations WHERE id = ?", (conversation_id,))
    conn.commit()
    conn.close()


def toggle_pin_conversation(conversation_id: int) -> bool:
    """Flips a conversation's pinned state, returns the new state."""
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT pinned FROM conversations WHERE id = ?", (conversation_id,)).fetchone()
    new_state = 0 if (row and row[0]) else 1
    conn.execute("UPDATE conversations SET pinned = ? WHERE id = ?", (new_state, conversation_id))
    conn.commit()
    conn.close()
    return bool(new_state)


def rename_project(project_id: int, new_name: str):
    """Renames a project's DISPLAY name only - never touches the actual
    folder on disk. The folder path stays exactly what it was; only how
    it's labeled in ORACLE's sidebar changes."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE projects SET name = ? WHERE id = ?", (new_name, project_id))
    conn.commit()
    conn.close()


def delete_project(project_id: int):
    """Removes a project from ORACLE's tracking - its conversation memory
    is deleted from ORACLE's database, but the actual project folder and
    every file in it are completely untouched. Re-opening the same folder
    later would simply start it fresh, as if new."""
    conn = sqlite3.connect(DB_PATH)
    conv_ids = [r[0] for r in conn.execute(
        "SELECT id FROM conversations WHERE project_id = ?", (project_id,)
    ).fetchall()]
    for conv_id in conv_ids:
        conn.execute("DELETE FROM messages WHERE conversation_id = ?", (conv_id,))
        conn.execute("DELETE FROM conversations WHERE id = ?", (conv_id,))
    conn.execute("DELETE FROM projects WHERE id = ?", (project_id,))
    conn.commit()
    conn.close()


def toggle_pin_project(project_id: int) -> bool:
    """Flips a project's pinned state, returns the new state."""
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute("SELECT pinned FROM projects WHERE id = ?", (project_id,)).fetchone()
    new_state = 0 if (row and row[0]) else 1
    conn.execute("UPDATE projects SET pinned = ? WHERE id = ?", (new_state, project_id))
    conn.commit()
    conn.close()
    return bool(new_state)


def load_conversation_messages(conversation_id: int) -> list:
    """Loads every message belonging to one conversation thread, in
    order - used when reopening a past chat from the sidebar. Unlike
    load_recent_history below, this isn't trimmed to a tail window, since
    a thread reopened from the sidebar should load in full."""
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT role, content, tool_calls, tool_call_id FROM messages "
        "WHERE conversation_id = ? ORDER BY id ASC",
        (conversation_id,),
    ).fetchall()
    conn.close()

    history = []
    for role, content, tool_calls_json, tool_call_id in rows:
        msg = {"role": role, "content": content}
        if tool_calls_json:
            msg["tool_calls"] = json.loads(tool_calls_json)
        if tool_call_id:
            msg["tool_call_id"] = tool_call_id
        history.append(msg)
    return history


def save_message(message: dict, conversation_id: int = None):
    """Appends a single message to the database, tagged to whichever
    conversation thread it belongs to, and bumps that conversation's
    updated_at so the sidebar sorts by most-recently-active."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        "INSERT INTO messages (conversation_id, role, content, tool_calls, tool_call_id, timestamp) "
        "VALUES (?, ?, ?, ?, ?, ?)",
        (
            conversation_id,
            message.get("role"),
            message.get("content"),
            # tool_calls is a list/dict - SQLite only stores text/numbers/blobs,
            # so we serialize it to a JSON string and deserialize on load.
            json.dumps(message["tool_calls"]) if message.get("tool_calls") else None,
            message.get("tool_call_id"),
            datetime.now().isoformat(),
        ),
    )
    if conversation_id is not None:
        conn.execute(
            "UPDATE conversations SET updated_at = ? WHERE id = ?",
            (datetime.now().isoformat(), conversation_id),
        )
    conn.commit()
    conn.close()


def load_recent_history(limit: int = MAX_HISTORY_MESSAGES) -> list:
    """Loads the most recent `limit` messages from disk, regardless of
    conversation thread - kept for the terminal test-mode fallback at the
    bottom of this file (`python core.py`), which doesn't have the concept
    of separate threads. The real GUI app uses load_conversation_messages
    instead."""
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT role, content, tool_calls, tool_call_id FROM messages "
        "ORDER BY id DESC LIMIT ?",
        (limit,),
    ).fetchall()
    conn.close()

    rows.reverse()  # we fetched newest-first; put back in chronological order

    history = []
    for role, content, tool_calls_json, tool_call_id in rows:
        msg = {"role": role, "content": content}
        if tool_calls_json:
            msg["tool_calls"] = json.loads(tool_calls_json)
        if tool_call_id:
            msg["tool_call_id"] = tool_call_id
        history.append(msg)

    # Same boundary issue as trim_history: if the LIMIT cut lands right after
    # an assistant tool_calls message but before its tool response, drop the
    # orphaned leading tool message(s).
    while history and history[0]["role"] == "tool":
        history.pop(0)

    return history


# ---------------------------------------------------------------------------
# STEP 1: Define a "tool" - a real Python function the LLM can decide to call.
# This is the foundation of item #6 (basic computer commands). The LLM never
# runs code itself - it just tells us "call get_current_time with these args"
# and WE execute it and hand the result back.
# ---------------------------------------------------------------------------

def get_current_time() -> str:
    """Returns the current date and time."""
    return datetime.now().strftime("%A, %B %d, %Y at %I:%M %p")


def list_files(directory: str = ".") -> str:
    """Lists files in a given directory."""
    try:
        files = os.listdir(directory)
    except Exception as e:
        return f"Error: {e}"

    if len(files) > MAX_LISTED_FILES:
        shown = json.dumps(files[:MAX_LISTED_FILES])
        return f"{shown}\n(showing {MAX_LISTED_FILES} of {len(files)} entries)"
    return json.dumps(files)


def open_file(path: str) -> str:
    """Opens a file with whatever application is set as the OS default for
    its type - e.g. a .pdf opens in your PDF viewer, a .docx in Word. This
    is exactly what happens when you double-click a file in File Explorer."""
    if not os.path.isfile(path):
        return f"Error: '{path}' does not exist or is not a file."

    try:
        if sys.platform == "win32":
            os.startfile(path)  # Windows-only; launches the default handler
        elif sys.platform == "darwin":
            subprocess.run(["open", path], check=True)
        else:
            subprocess.run(["xdg-open", path], check=True)
        return f"Opened '{path}'."
    except Exception as e:
        return f"Error opening '{path}': {e}"


def _find_start_menu_shortcut(app_name: str):
    """
    Searches Windows Start Menu shortcut folders for a .lnk file whose name
    contains app_name (case-insensitive) - this is essentially what happens
    when you press the Windows key and type an app's name. Most installed
    apps (Obsidian included) aren't on PATH at all; they only exist as a
    shortcut here, which is why a plain PATH-based launch fails for them.
    Returns the shortcut's full path if found, else None.
    """
    search_dirs = []
    appdata = os.environ.get("APPDATA")
    programdata = os.environ.get("PROGRAMDATA")
    if appdata:
        search_dirs.append(os.path.join(appdata, "Microsoft", "Windows", "Start Menu", "Programs"))
    if programdata:
        search_dirs.append(os.path.join(programdata, "Microsoft", "Windows", "Start Menu", "Programs"))

    name_lower = app_name.lower()
    for base in search_dirs:
        if not os.path.isdir(base):
            continue
        for root, _, files in os.walk(base):
            for f in files:
                if f.lower().endswith(".lnk") and name_lower in f.lower():
                    return os.path.join(root, f)
    return None


def launch_app(app_name: str) -> str:
    """
    Launches an application by name (e.g. "notepad", "chrome", "obsidian")
    WITHOUT searching the general filesystem first. Two strategies, tried
    in order:

      1. Look for a matching Start Menu shortcut (.lnk) and launch that
         directly - this covers the vast majority of installed apps, since
         that's genuinely how Windows itself finds them by name.
      2. Fall back to letting the shell resolve the name via PATH - this
         covers built-in commands like "notepad" or "calc" that don't have
         Start Menu shortcuts but ARE directly runnable.

    Unlike a naive Popen-and-forget, this actually checks whether the
    launch succeeded: a failing command (like "not recognized") normally
    exits almost immediately, so we wait briefly and inspect the result
    instead of always reporting success.
    """
    if sys.platform == "win32":
        shortcut = _find_start_menu_shortcut(app_name)
        if shortcut:
            try:
                os.startfile(shortcut)
                return f"Launched '{app_name}' (via Start Menu shortcut)."
            except Exception as e:
                return f"Error launching '{app_name}': {e}"

        # No shortcut found - fall back to PATH resolution via the shell.
        try:
            proc = subprocess.Popen(
                app_name, shell=True,
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True,
            )
            try:
                # A command that fails to resolve (like "not recognized")
                # exits almost instantly. A real GUI app keeps running, so
                # this timeout is what lets us tell the two apart without
                # blocking forever on a successful, long-running launch.
                stdout, stderr = proc.communicate(timeout=1.5)
                if proc.returncode != 0:
                    detail = stderr.strip() or stdout.strip() or f"exit code {proc.returncode}"
                    return (
                        f"Error launching '{app_name}': {detail}. "
                        f"No Start Menu shortcut matched this name either - "
                        f"try the app's exact display name."
                    )
                return f"Launched '{app_name}'."
            except subprocess.TimeoutExpired:
                # Still running past the timeout - treat as a successful,
                # ongoing GUI app launch rather than waiting indefinitely.
                return f"Launched '{app_name}'."
        except Exception as e:
            return f"Error launching '{app_name}': {e}"

    elif sys.platform == "darwin":
        try:
            subprocess.run(["open", "-a", app_name], check=True)
            return f"Launched '{app_name}'."
        except Exception as e:
            return f"Error launching '{app_name}': {e}"
    else:
        try:
            subprocess.Popen([app_name])
            return f"Launched '{app_name}'."
        except Exception as e:
            return f"Error launching '{app_name}': {e}"


# How many characters of extracted file text we hand to the model at once.
# Same reasoning as MAX_LISTED_FILES: a 40-page PDF's full text would blow
# past the context window and get resent on every later turn.
MAX_FILE_CHARS = 6000


def read_file(path: str) -> str:
    """Extracts and returns text content from a file so ORACLE can summarize
    it or answer questions about it. Supports .txt/.md, .pdf, and .docx -
    the format is picked automatically from the file extension."""
    if not os.path.isfile(path):
        return f"Error: '{path}' does not exist or is not a file."

    ext = os.path.splitext(path)[1].lower()

    try:
        if ext in (".txt", ".md", ".csv", ".json", ".log"):
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()

        elif ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)

        elif ext == ".docx":
            from docx import Document
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)

        else:
            return f"Error: unsupported file type '{ext}'. Supported: .txt, .md, .csv, .json, .log, .pdf, .docx"

    except Exception as e:
        return f"Error reading '{path}': {e}"

    if not text.strip():
        return f"'{path}' appears to be empty or its text couldn't be extracted (e.g. a scanned/image-only PDF)."

    if len(text) > MAX_FILE_CHARS:
        return f"{text[:MAX_FILE_CHARS]}\n\n(truncated - showing first {MAX_FILE_CHARS} of {len(text)} characters)"
    return text


def move_file(source: str, destination: str) -> str:
    """Moves or renames a file. If destination is a directory, the file is
    moved into it keeping its original name; otherwise destination is
    treated as the new full path/name."""
    if not os.path.isfile(source):
        return f"Error: source '{source}' does not exist or is not a file."
    try:
        shutil.move(source, destination)
        return f"Moved '{source}' to '{destination}'."
    except Exception as e:
        return f"Error moving '{source}' to '{destination}': {e}"


def delete_file(path: str) -> str:
    """Permanently deletes a file. There is no undo - this does not use the
    Recycle Bin, it removes the file directly."""
    if not os.path.isfile(path):
        return f"Error: '{path}' does not exist or is not a file."
    try:
        os.remove(path)
        return f"Deleted '{path}'."
    except Exception as e:
        return f"Error deleting '{path}': {e}"


def create_file(path: str, content: str = "") -> str:
    """Creates a new text file with the given content. Fails if the file
    already exists, to avoid silently overwriting something."""
    if os.path.exists(path):
        return f"Error: '{path}' already exists. Use a different name or delete it first."
    try:
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"Created '{path}'."
    except Exception as e:
        return f"Error creating '{path}': {e}"


def get_system_stats_dict() -> dict:
    """
    Returns live system stats as structured data (not a string) - used
    directly by the UI for the sidebar's live bars, and also as the source
    of truth behind the get_system_info tool below. Keeping this separate
    from the tool-facing text version means the UI can poll it every few
    seconds without spending any LLM tokens or API calls.
    """
    cpu_percent = psutil.cpu_percent(interval=0.2)
    mem = psutil.virtual_memory()
    disk_path = "C:\\" if sys.platform == "win32" else "/"
    disk = psutil.disk_usage(disk_path)

    stats = {
        "cpu_percent": round(cpu_percent, 1),
        "ram_percent": round(mem.percent, 1),
        "ram_used_gb": round(mem.used / (1024 ** 3), 1),
        "ram_total_gb": round(mem.total / (1024 ** 3), 1),
        "disk_percent": round(disk.percent, 1),
        "disk_used_gb": round(disk.used / (1024 ** 3), 1),
        "disk_total_gb": round(disk.total / (1024 ** 3), 1),
        "gpu": None,
    }

    # VRAM: only meaningful for a dedicated GPU. GPUtil supports NVIDIA
    # cards; on integrated-graphics machines (or without GPUtil/no NVIDIA
    # driver) this will fail, and we deliberately leave "gpu" as None
    # rather than fabricate a number.
    try:
        import GPUtil
        gpus = GPUtil.getGPUs()
        if gpus:
            g = gpus[0]
            stats["gpu"] = {
                "name": g.name,
                "vram_percent": round(g.memoryUtil * 100, 1),
                "vram_used_mb": round(g.memoryUsed),
                "vram_total_mb": round(g.memoryTotal),
            }
    except Exception:
        pass

    return stats


def get_system_info() -> str:
    """Tool-facing wrapper: formats live system stats as text for ORACLE
    to read out or discuss in conversation."""
    s = get_system_stats_dict()
    lines = [
        f"CPU usage: {s['cpu_percent']}%",
        f"RAM usage: {s['ram_percent']}% ({s['ram_used_gb']} GB of {s['ram_total_gb']} GB)",
        f"Storage usage: {s['disk_percent']}% ({s['disk_used_gb']} GB of {s['disk_total_gb']} GB)",
    ]
    if s["gpu"]:
        g = s["gpu"]
        lines.append(
            f"GPU: {g['name']} - VRAM {g['vram_percent']}% "
            f"({g['vram_used_mb']} MB of {g['vram_total_mb']} MB)"
        )
    else:
        lines.append("GPU/VRAM: no dedicated GPU detected (integrated graphics).")
    return "\n".join(lines)


# Lazily created so a missing TAVILY_API_KEY doesn't crash the whole app on
# startup - it only becomes a problem the moment web_search is actually
# called, with a clear error message instead of a startup traceback.
_tavily_client = None


def web_search(query: str) -> str:
    """Searches the web via Tavily and returns a handful of results with
    titles, URLs, and short content summaries."""
    global _tavily_client

    api_key = os.environ.get("TAVILY_API_KEY")
    if not api_key:
        return "Error: TAVILY_API_KEY environment variable is not set."

    if _tavily_client is None:
        _tavily_client = TavilyClient(api_key=api_key)

    try:
        response = _tavily_client.search(query=query, max_results=5)
    except Exception as e:
        return f"Error performing web search: {e}"

    results = response.get("results", [])
    if not results:
        return "No results found."

    formatted = []
    for r in results:
        title = r.get("title", "Untitled")
        url = r.get("url", "")
        content = (r.get("content") or "")[:400]
        formatted.append(f"{title}\n{url}\n{content}")

    return "\n\n".join(formatted)


# ---------------------------------------------------------------------------
# CODING AGENT: a scoped, single-tool version of "an agent for coding" -
# not a separate multi-agent framework, just ORACLE's main model
# delegating coding-heavy requests to a coding-specialist model when it
# judges that's a better fit than answering directly itself.
#
# Uses Qwen3.6-27B (qwen/qwen3.6-27b) - same Groq client, same free
# rate-limited developer tier, no new account or setup needed. This was
# originally Kimi K2, but Groq fully deprecated that model (twice - first
# the original, then its successor) as of April 2026, so it's no longer
# available at all. Worth knowing: Groq retires and replaces models on an
# ongoing basis, so this model ID may itself need updating again someday -
# check https://console.groq.com/docs/deprecations if this ever starts
# erroring with a 404.
# ---------------------------------------------------------------------------

CODING_MODEL = "qwen/qwen3.6-27b"
CODING_AGENT_SYSTEM_PROMPT = (
    "You are an expert coding assistant. Provide clean, correct, well-explained "
    "code and precise technical answers to programming questions. Explain your "
    "reasoning where it genuinely helps understanding, but don't pad answers with "
    "unnecessary chatter - the person asking wants a working answer."
)


def ask_coding_agent(prompt: str) -> str:
    """
    Delegates a coding-focused question or task to a coding-specialist model
    (Qwen3.6-27B) rather than answering it directly - use this for non-trivial
    coding help: writing code, debugging, explaining code, architecture
    questions. Pass the specific coding question/task as the prompt; this
    doesn't have access to the rest of the conversation, so include
    whatever context (language, relevant code, error messages) is actually
    needed to answer well.
    """
    try:
        response = client.chat.completions.create(
            model=CODING_MODEL,
            messages=[
                {"role": "system", "content": CODING_AGENT_SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error consulting coding agent: {e}"


# ---------------------------------------------------------------------------
# PROJECTS: a separate, restricted tool-calling loop for the coding agent
# when working within a specific project folder. Unlike ask_coding_agent
# above (stateless, no tools, no file access), this gives it real
# read/write access to files - but ONLY within the selected project's own
# folder. Path-traversal protection (_safe_project_path) is the important
# part here: every file tool call is validated to resolve to somewhere
# genuinely inside the project root before touching disk, so a
# hallucinated or malicious path like "../../../Desktop/important.txt"
# gets refused rather than silently escaping the sandbox.
# ---------------------------------------------------------------------------

PROJECT_TOOLS_SCHEMA = [
    {
        "type": "function",
        "function": {
            "name": "list_project_files",
            "description": "List files and folders within the current project (or a subdirectory of it).",
            "parameters": {
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "Subdirectory within the project to list, relative to the project root. Defaults to the project root itself.",
                    }
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_project_file",
            "description": "Read the contents of a file within the current project.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the file, relative to the project root."}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_project_file",
            "description": (
                "Create or overwrite a file within the current project with the given "
                "content. Unlike the general assistant's create_file, this CAN overwrite "
                "an existing file - normal for iterative coding work."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the file, relative to the project root."},
                    "content": {"type": "string", "description": "Full content to write to the file."},
                },
                "required": ["path", "content"],
            },
        },
    },
]


def _safe_project_path(project_root: str, relative_path: str) -> str:
    """
    Resolves relative_path against project_root and verifies the result
    is genuinely still inside project_root, raising if not. This is the
    entire safety boundary for the project coding agent's file access -
    every read/write goes through this first.
    """
    root_abs = os.path.abspath(project_root)
    target_abs = os.path.abspath(os.path.join(root_abs, relative_path or "."))
    if target_abs != root_abs and not target_abs.startswith(root_abs + os.sep):
        raise ValueError(
            f"'{relative_path}' resolves outside the project folder - refused for safety."
        )
    return target_abs


def _make_project_tools(project_root: str) -> dict:
    """Builds the three project-scoped tool functions as closures over a
    specific project_root, so each project conversation gets its own
    correctly-sandboxed set - these are never shared globally like the
    main AVAILABLE_FUNCTIONS dict."""

    def project_list_files(directory: str = ".") -> str:
        try:
            safe_dir = _safe_project_path(project_root, directory)
        except ValueError as e:
            return f"Error: {e}"
        return list_files(safe_dir)

    def project_read_file(path: str) -> str:
        try:
            safe_path = _safe_project_path(project_root, path)
        except ValueError as e:
            return f"Error: {e}"

        if not os.path.isfile(safe_path):
            return f"Error: '{path}' does not exist or is not a file."

        ext = os.path.splitext(safe_path)[1].lower()
        try:
            if ext == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(safe_path)
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            elif ext == ".docx":
                from docx import Document
                doc = Document(safe_path)
                text = "\n".join(p.text for p in doc.paragraphs)
            else:
                # Any other extension - source code, config, markup,
                # whatever - reads as plain text. Unlike the general
                # assistant's read_file (which whitelists a handful of
                # document types), project files are overwhelmingly
                # source code, so this needs to be permissive rather
                # than restrictive.
                with open(safe_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
        except Exception as e:
            return f"Error reading '{path}': {e}"

        if not text.strip():
            return f"'{path}' appears to be empty."

        if len(text) > MAX_FILE_CHARS:
            return f"{text[:MAX_FILE_CHARS]}\n\n(truncated - showing first {MAX_FILE_CHARS} of {len(text)} characters)"
        return text

    def project_write_file(path: str, content: str) -> str:
        try:
            safe_path = _safe_project_path(project_root, path)
        except ValueError as e:
            return f"Error: {e}"
        try:
            parent = os.path.dirname(safe_path)
            if parent:
                os.makedirs(parent, exist_ok=True)
            with open(safe_path, "w", encoding="utf-8") as f:
                f.write(content)
            return f"Wrote '{path}'."
        except Exception as e:
            return f"Error writing '{path}': {e}"

    return {
        "list_project_files": project_list_files,
        "read_project_file": project_read_file,
        "write_project_file": project_write_file,
    }


def run_project_conversation(user_input: str, history: list, conversation_id: int, project_root: str) -> str:
    """
    Like run_conversation, but scoped to one project: uses the coding
    specialist model, the restricted file-only tool set above (sandboxed
    to project_root), and its own conversation thread - completely
    isolated from Main Chat and every other project's memory.
    """
    project_tools = _make_project_tools(project_root)

    history[:] = trim_history(history)
    user_msg = {"role": "user", "content": user_input}
    history.append(user_msg)
    save_message(user_msg, conversation_id)

    for _ in range(MAX_TOOL_ROUNDS):
        response = client.chat.completions.create(
            model=CODING_MODEL,
            messages=history,
            tools=PROJECT_TOOLS_SCHEMA,
        )
        message = response.choices[0].message

        if not message.tool_calls:
            final_msg = {"role": "assistant", "content": message.content}
            history.append(final_msg)
            save_message(final_msg, conversation_id)
            return message.content

        tool_call_msg = {
            "role": "assistant",
            "content": message.content,
            "tool_calls": [
                {
                    "id": tc.id,
                    "type": "function",
                    "function": {"name": tc.function.name, "arguments": tc.function.arguments},
                }
                for tc in message.tool_calls
            ],
        }
        history.append(tool_call_msg)
        save_message(tool_call_msg, conversation_id)

        for tool_call in message.tool_calls:
            fn_name = tool_call.function.name
            fn_args = json.loads(tool_call.function.arguments)
            fn = project_tools.get(fn_name)
            result = fn(**fn_args) if fn else f"Unknown tool: {fn_name}"
            tool_result_msg = {
                "role": "tool",
                "tool_call_id": tool_call.id,
                "content": str(result),
            }
            history.append(tool_result_msg)
            save_message(tool_result_msg, conversation_id)

    give_up = f"I kept calling tools without reaching an answer ({MAX_TOOL_ROUNDS} rounds)."
    give_up_msg = {"role": "assistant", "content": give_up}
    history.append(give_up_msg)
    save_message(give_up_msg, conversation_id)
    return give_up


def project_system_prompt(project_name: str) -> str:
    """System prompt for a project conversation - same coding-specialist
    persona, plus explicit awareness of which project it's scoped to."""
    return (
        CODING_AGENT_SYSTEM_PROMPT
        + f" You are currently working within the '{project_name}' project folder. Use "
        "list_project_files, read_project_file, and write_project_file to explore and "
        "modify files there - these tools are sandboxed to this project only and cannot "
        "read or write anything outside it, even if asked."
    )


def open_web_search(query: str) -> str:
    """
    Opens the user's default web browser to a real, visible search results
    page - unlike web_search above, which fetches results as text for
    ORACLE to read and summarize, this actually launches a browser tab so
    the user can look at and interact with the results themselves.
    """
    url = f"https://www.google.com/search?q={quote_plus(query)}"
    try:
        webbrowser.open(url, new=2)  # new=2: open in a new tab, not the same window
        return f"Opened a web search for '{query}' in your browser."
    except Exception as e:
        return f"Error opening web search: {e}"


def open_url(url: str) -> str:
    """
    Opens a specific web address in the user's default browser - for when
    the user names an actual site to go to (e.g. "open youtube.com"),
    as opposed to searching for something (see open_web_search/web_search).
    """
    url = url.strip()

    # Add a scheme if the user just said "youtube.com" without one.
    if not re.match(r"^https?://", url, re.IGNORECASE):
        url = f"https://{url}"

    # Basic sanity check - must look like an actual domain, and only
    # http(s) is allowed (guards against something like a javascript: or
    # file: URI slipping through, which "open this website" should never
    # trigger).
    if not re.match(r"^https?://[^\s]+\.[^\s]+", url, re.IGNORECASE):
        return f"Error: '{url}' doesn't look like a valid web address."

    try:
        webbrowser.open(url, new=2)
        return f"Opened {url}"
    except Exception as e:
        return f"Error opening '{url}': {e}"


# ---------------------------------------------------------------------------
# VOICE INPUT: local, free speech-to-text via faster-whisper. Recording
# happens entirely in Python (via sounddevice) rather than in the browser -
# this keeps audio capture and transcription in one place instead of
# encoding audio in JS and shipping it across the pywebview bridge.
# ---------------------------------------------------------------------------

# "base" balances speed and accuracy reasonably well on CPU-only hardware
# (no dedicated GPU here). "tiny" is faster but noticeably less accurate;
# "small"/"medium" are more accurate but slower per transcription.
WHISPER_MODEL_SIZE = "base"
SAMPLE_RATE = 16000  # Whisper's native input rate - recording at this rate
                      # directly avoids a separate resampling step.

_whisper_model = None
_recording_stream = None
_recording_frames = []


def _get_whisper_model():
    """Lazily loads the Whisper model on first use (this can take a few
    seconds and downloads model weights on first run ever) rather than
    slowing down every ORACLE startup, most of which won't use voice."""
    global _whisper_model
    if _whisper_model is None:
        from faster_whisper import WhisperModel
        _whisper_model = WhisperModel(WHISPER_MODEL_SIZE, device="cpu", compute_type="int8")
    return _whisper_model


def start_recording() -> str:
    """Begins capturing microphone audio into memory. Call
    stop_recording_and_transcribe() to end it and get the transcribed text."""
    global _recording_stream, _recording_frames

    if _recording_stream is not None:
        return "Already recording."

    _recording_frames = []

    def _callback(indata, frames, time_info, status):
        _recording_frames.append(indata.copy())

    try:
        _recording_stream = sd.InputStream(
            samplerate=SAMPLE_RATE, channels=1, dtype="float32", callback=_callback
        )
        _recording_stream.start()
        return "Recording started."
    except Exception as e:
        _recording_stream = None
        return f"Error starting recording: {e}"


def stop_recording_and_transcribe() -> str:
    """Stops capturing audio and transcribes whatever was recorded."""
    global _recording_stream, _recording_frames

    if _recording_stream is None:
        return "Error: not currently recording."

    try:
        _recording_stream.stop()
        _recording_stream.close()
    except Exception as e:
        _recording_stream = None
        return f"Error stopping recording: {e}"

    _recording_stream = None

    if not _recording_frames:
        return "Error: no audio was captured."

    audio = np.concatenate(_recording_frames, axis=0).flatten()
    _recording_frames = []

    # Rough silence check - a very quiet/empty recording (e.g. mic muted,
    # or the stop button hit almost immediately) shouldn't be sent to
    # Whisper at all, since it tends to hallucinate text from near-silence.
    if np.abs(audio).mean() < 0.001:
        return "Error: no speech detected (audio was silent)."

    try:
        model = _get_whisper_model()
        segments, _ = model.transcribe(audio, language=None)
        text = " ".join(segment.text.strip() for segment in segments).strip()
    except Exception as e:
        return f"Error transcribing audio: {e}"

    return text if text else "Error: could not understand the audio."


# Tunable VAD (voice activity detection) parameters for listen_and_transcribe.
CALIBRATION_DURATION_SEC = 0.5   # brief ambient-noise measurement before listening starts
SPEECH_THRESHOLD_MULTIPLIER = 3.5  # how far above the measured noise floor counts as speech
MIN_SPEECH_THRESHOLD = 0.01      # floor, so a near-silent room doesn't make the mic's own
                                  # self-noise register as "speech"
SILENCE_DURATION_SEC = 1.2       # pause length (after speech) that ends listening
NO_SPEECH_TIMEOUT_SEC = 8.0      # give up if nothing is said within this long (calibration
                                  # eats into this window, so it's a bit longer than before)
MAX_RECORDING_SEC = 20.0         # hard safety cap regardless of VAD state


def listen_and_transcribe() -> str:
    """
    Listens via the microphone until you pause after speaking (or a safety
    timeout elapses), then transcribes what was captured - a single
    blocking call, unlike start_recording/stop_recording_and_transcribe's
    manual two-step. This is what the global hotkey uses: press it, talk,
    stop talking, and this returns once it detects you're done - the same
    turn-taking style used by voice mode in other AI chat apps.

    The first CALIBRATION_DURATION_SEC of audio is used to measure YOUR
    mic's actual ambient noise floor rather than assuming a fixed volume
    level - different microphones and rooms have very different baseline
    noise, so a single hardcoded threshold either misses quiet speech or
    (more often) never sees true silence at all if the room's baseline is
    louder than the hardcoded guess.
    """
    frames = []
    done = threading.Event()
    state = {
        "speech_detected": False,
        "silence_frames": 0,
        "total_frames": 0,
        "calibration_rms": [],
        "threshold": None,  # set once calibration finishes
    }

    def callback(indata, frame_count, time_info, status):
        frames.append(indata.copy())
        state["total_frames"] += frame_count
        rms = float(np.sqrt(np.mean(indata.astype(np.float64) ** 2)))
        elapsed = state["total_frames"] / SAMPLE_RATE

        # Phase 1: calibration - just measure, don't judge speech/silence yet.
        if state["threshold"] is None:
            state["calibration_rms"].append(rms)
            if elapsed >= CALIBRATION_DURATION_SEC:
                noise_floor = float(np.mean(state["calibration_rms"]))
                state["threshold"] = max(
                    noise_floor * SPEECH_THRESHOLD_MULTIPLIER, MIN_SPEECH_THRESHOLD
                )
            return

        # Phase 2: normal VAD logic, using the calibrated threshold.
        if rms > state["threshold"]:
            state["speech_detected"] = True
            state["silence_frames"] = 0
        elif state["speech_detected"]:
            state["silence_frames"] += frame_count

        silence_elapsed = state["silence_frames"] / SAMPLE_RATE

        if state["speech_detected"] and silence_elapsed >= SILENCE_DURATION_SEC:
            raise sd.CallbackStop()
        if not state["speech_detected"] and elapsed >= NO_SPEECH_TIMEOUT_SEC:
            raise sd.CallbackStop()
        if elapsed >= MAX_RECORDING_SEC:
            raise sd.CallbackStop()

    def _finished():
        done.set()

    try:
        stream = sd.InputStream(
            samplerate=SAMPLE_RATE,
            channels=1,
            dtype="float32",
            blocksize=int(SAMPLE_RATE * 0.1),  # ~100ms chunks for responsive VAD
            callback=callback,
            finished_callback=_finished,
        )
    except Exception as e:
        return f"Error starting microphone: {e}"

    with stream:
        # Extra couple seconds of margin beyond MAX_RECORDING_SEC, purely
        # so this wait() can't itself hang forever in some edge case.
        done.wait(timeout=MAX_RECORDING_SEC + 2)

    if not frames:
        return "Error: no audio was captured."

    audio = np.concatenate(frames, axis=0).flatten()

    if not state["speech_detected"]:
        return "Error: no speech detected."

    try:
        model = _get_whisper_model()
        segments, _ = model.transcribe(audio, language=None)
        text = " ".join(segment.text.strip() for segment in segments).strip()
    except Exception as e:
        return f"Error transcribing audio: {e}"

    return text if text else "Error: could not understand the audio."



# ---------------------------------------------------------------------------
# TEXT-TO-SPEECH: local, free voice output via Piper. Every ORACLE reply
# gets spoken, whether it came from typing or voice - see ui.py for where
# this actually gets called (kept out of the conversational logic here so
# core.py doesn't need to know whether a reply is about to hit a chat
# window, a voice turn, or both).
# ---------------------------------------------------------------------------

PIPER_VOICE_NAME = "en_GB-alan-medium"  # smooth British male voice - fits ORACLE's character
_piper_voice = None


def _strip_markdown_for_speech(text: str) -> str:
    """
    Removes markdown formatting before handing text to Piper - without
    this, a reply like "**Note:** run `python ui.py`" would be read aloud
    literally as "asterisk asterisk Note asterisk asterisk colon run
    backtick python ui dot py backtick", which is unusable. Code blocks
    are dropped entirely rather than read aloud, since spoken code is
    rarely useful and often very long.
    """
    text = re.sub(r"```.*?```", "", text, flags=re.DOTALL)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^[-*•]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
    return text.strip()


def _get_piper_voice():
    """
    Lazily loads (and downloads, on first-ever use) the Piper voice model.
    Downloaded once into ORACLE's own data directory - after that first
    run, this is instant and fully offline.
    """
    global _piper_voice
    if _piper_voice is None:
        from piper import PiperVoice
        from piper.download_voices import download_voice

        voice_dir = Path(_get_data_dir()) / "voices"
        voice_dir.mkdir(parents=True, exist_ok=True)

        download_voice(PIPER_VOICE_NAME, voice_dir)  # no-ops if already downloaded

        model_path = voice_dir / f"{PIPER_VOICE_NAME}.onnx"
        config_path = voice_dir / f"{PIPER_VOICE_NAME}.onnx.json"
        _piper_voice = PiperVoice.load(str(model_path), str(config_path))
    return _piper_voice


def speak(text: str) -> str:
    """
    Synthesizes text to speech locally via Piper and plays it through the
    default audio output. Blocks until playback finishes - callers that
    don't want to wait (e.g. so a chat message can appear immediately
    instead of only after the audio finishes) should run this in a
    background thread instead of calling it directly.
    """
    cleaned = _strip_markdown_for_speech(text)
    if not cleaned:
        return "Error: no speakable text after removing formatting."

    try:
        voice = _get_piper_voice()
        chunks = list(voice.synthesize(cleaned))
        if not chunks:
            return "Error: no audio was generated."

        audio = np.concatenate([c.audio_float_array for c in chunks])
        sample_rate = chunks[0].sample_rate

        sd.play(audio, samplerate=sample_rate)
        sd.wait()
        return "Spoken."
    except Exception as e:
        return f"Error speaking: {e}"


# Lazily created for the same reason as the Tavily client - a missing or
# broken toast library shouldn't crash the whole app at startup, only the
# one tool that actually needs it.
_toaster = None


def send_notification(title: str, message: str = "") -> str:
    """
    Shows a real Windows toast notification (the kind that pops up from
    the notification area), separate from anything shown in the ORACLE
    chat window itself - useful for things ORACLE wants to flag even if
    you're not actively looking at the app.
    """
    if sys.platform != "win32":
        return "Error: toast notifications are only supported on Windows."

    global _toaster
    try:
        from windows_toasts import Toast, WindowsToaster

        if _toaster is None:
            _toaster = WindowsToaster("ORACLE")

        toast = Toast()
        toast.text_fields = [title, message] if message else [title]
        _toaster.show_toast(toast)
        return f"Notification sent: '{title}'"
    except Exception as e:
        return f"Error sending notification: {e}"


# ---------------------------------------------------------------------------
# OUTLOOK (Microsoft Graph): Mail and Calendar share one login, since
# Graph is Microsoft's unified API for both. Uses MSAL's "public client"
# flow - no client secret embedded in the app (unlike Google's typical
# desktop-app OAuth pattern), which is the more secure option for
# something distributed as a standalone .exe.
#
# SETUP (one-time, on your end - I can't do this part for you):
#   1. Go to https://portal.azure.com -> Microsoft Entra ID -> App registrations
#   2. New registration - name it whatever, "Personal" account types is fine
#   3. No redirect URI needed for the interactive flow used here
#   4. Copy the "Application (client) ID" and set it as an environment
#      variable: MS_CLIENT_ID
#   5. Under "API permissions", add: Mail.ReadWrite, Mail.Send,
#      Calendars.ReadWrite (delegated permissions, not application)
#
# First run opens your browser for a one-time login/consent. After that,
# the token is cached to disk and silently refreshed - no repeated logins.
# ---------------------------------------------------------------------------

GRAPH_SCOPES = ["Mail.ReadWrite", "Mail.Send", "Calendars.ReadWrite"]
GRAPH_BASE = "https://graph.microsoft.com/v1.0"

_msal_app = None
_token_cache = None


def _get_token_cache_path() -> str:
    return os.path.join(_get_data_dir(), "ms_token_cache.bin")


def _get_graph_token() -> str:
    """
    Returns a valid Microsoft Graph access token, handling the whole MSAL
    dance: load any cached session from disk, try to reuse it silently
    (no user interaction), and only fall back to opening a browser for a
    fresh login if there's no valid cached session at all.
    """
    global _msal_app, _token_cache

    client_id = os.environ.get("MS_CLIENT_ID")
    if not client_id:
        raise RuntimeError("MS_CLIENT_ID environment variable is not set.")

    if _token_cache is None:
        _token_cache = msal.SerializableTokenCache()
        cache_path = _get_token_cache_path()
        if os.path.exists(cache_path):
            with open(cache_path, "r") as f:
                _token_cache.deserialize(f.read())

    if _msal_app is None:
        _msal_app = msal.PublicClientApplication(client_id, token_cache=_token_cache)

    result = None
    accounts = _msal_app.get_accounts()
    if accounts:
        result = _msal_app.acquire_token_silent(GRAPH_SCOPES, account=accounts[0])

    if not result:
        result = _msal_app.acquire_token_interactive(GRAPH_SCOPES)

    if _token_cache.has_state_changed:
        with open(_get_token_cache_path(), "w") as f:
            f.write(_token_cache.serialize())

    if "access_token" not in result:
        error_desc = result.get("error_description", "unknown error")
        raise RuntimeError(f"Microsoft sign-in failed: {error_desc}")

    return result["access_token"]


def _graph_request(method: str, path: str, json_body: dict = None) -> requests.Response:
    """Thin wrapper around a Graph API call - attaches the bearer token
    and the base URL so individual tools below don't repeat that setup."""
    token = _get_graph_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    return requests.request(method, f"{GRAPH_BASE}{path}", headers=headers, json=json_body, timeout=15)


def list_recent_emails(count: int = 10) -> str:
    """Lists the most recent emails in the Outlook inbox."""
    try:
        resp = _graph_request(
            "GET",
            f"/me/messages?$top={count}&$select=id,subject,from,receivedDateTime,bodyPreview&$orderby=receivedDateTime desc",
        )
        resp.raise_for_status()
    except Exception as e:
        return f"Error listing emails: {e}"

    messages = resp.json().get("value", [])
    if not messages:
        return "No emails found."

    lines = []
    for m in messages:
        sender = m.get("from", {}).get("emailAddress", {}).get("address", "unknown sender")
        lines.append(
            f"ID: {m['id']}\nFrom: {sender}\nSubject: {m.get('subject', '(no subject)')}\n"
            f"Received: {m.get('receivedDateTime', '')}\nPreview: {m.get('bodyPreview', '')[:200]}"
        )
    return "\n\n".join(lines)


def read_email(email_id: str) -> str:
    """Reads the full content of a specific email by its ID (get the ID
    from list_recent_emails first)."""
    try:
        resp = _graph_request("GET", f"/me/messages/{email_id}?$select=subject,from,body,receivedDateTime")
        resp.raise_for_status()
    except Exception as e:
        return f"Error reading email: {e}"

    m = resp.json()
    sender = m.get("from", {}).get("emailAddress", {}).get("address", "unknown sender")
    body = m.get("body", {}).get("content", "")
    return f"From: {sender}\nSubject: {m.get('subject', '(no subject)')}\n\n{body}"


def send_email(to: str, subject: str, body: str) -> str:
    """Sends an email from the user's Outlook account. This sends
    immediately - there is no draft/confirmation step."""
    payload = {
        "message": {
            "subject": subject,
            "body": {"contentType": "Text", "content": body},
            "toRecipients": [{"emailAddress": {"address": to}}],
        }
    }
    try:
        resp = _graph_request("POST", "/me/sendMail", json_body=payload)
        resp.raise_for_status()
        return f"Email sent to {to}."
    except Exception as e:
        return f"Error sending email: {e}"


def delete_email(email_id: str) -> str:
    """Deletes an email by ID. Microsoft moves deleted messages to the
    Deleted Items folder rather than erasing them immediately, so this is
    usually recoverable for a while - but treat it as final."""
    try:
        resp = _graph_request("DELETE", f"/me/messages/{email_id}")
        resp.raise_for_status()
        return f"Deleted email {email_id}."
    except Exception as e:
        return f"Error deleting email: {e}"


def list_upcoming_events(days: int = 7) -> str:
    """Lists upcoming Outlook calendar events over the next N days."""
    start = datetime.utcnow().isoformat() + "Z"
    end = (datetime.utcnow() + timedelta(days=days)).isoformat() + "Z"
    try:
        resp = _graph_request(
            "GET",
            f"/me/calendarView?startDateTime={start}&endDateTime={end}"
            f"&$select=id,subject,start,end,location&$orderby=start/dateTime",
        )
        resp.raise_for_status()
    except Exception as e:
        return f"Error listing events: {e}"

    events = resp.json().get("value", [])
    if not events:
        return f"No events in the next {days} day(s)."

    lines = []
    for e in events:
        loc = e.get("location", {}).get("displayName", "")
        lines.append(
            f"ID: {e['id']}\n{e.get('subject', '(no subject)')}\n"
            f"Start: {e.get('start', {}).get('dateTime', '')}\n"
            f"End: {e.get('end', {}).get('dateTime', '')}" + (f"\nLocation: {loc}" if loc else "")
        )
    return "\n\n".join(lines)


def create_calendar_event(subject: str, start_iso: str, end_iso: str, attendees: str = "") -> str:
    """
    Creates an Outlook calendar event. start_iso/end_iso must be ISO 8601
    datetimes (e.g. '2026-09-05T14:00:00'). attendees is an optional
    comma-separated list of email addresses.
    """
    payload = {
        "subject": subject,
        "start": {"dateTime": start_iso, "timeZone": "UTC"},
        "end": {"dateTime": end_iso, "timeZone": "UTC"},
    }
    if attendees:
        payload["attendees"] = [
            {"emailAddress": {"address": addr.strip()}, "type": "required"}
            for addr in attendees.split(",") if addr.strip()
        ]
    try:
        resp = _graph_request("POST", "/me/events", json_body=payload)
        resp.raise_for_status()
        return f"Created event '{subject}'."
    except Exception as e:
        return f"Error creating event: {e}"


def delete_calendar_event(event_id: str) -> str:
    """Deletes/cancels an Outlook calendar event by ID. This cannot be
    undone - attendees (if any) are sent a cancellation."""
    try:
        resp = _graph_request("DELETE", f"/me/events/{event_id}")
        resp.raise_for_status()
        return f"Deleted event {event_id}."
    except Exception as e:
        return f"Error deleting event: {e}"


# ---------------------------------------------------------------------------
# GMAIL: separate OAuth flow from Outlook - Google's, not Microsoft's.
# Uses google-auth-oauthlib's InstalledAppFlow, which briefly runs a local
# web server on your machine purely to catch the OAuth redirect after you
# approve access in your browser - normal for this kind of desktop app
# flow, nothing stays listening afterward.
#
# SETUP (one-time, on your end):
#   1. Go to https://console.cloud.google.com -> create a project (or
#      reuse one) -> APIs & Services -> Library -> enable "Gmail API"
#   2. APIs & Services -> Credentials -> Create Credentials -> OAuth
#      client ID -> Application type: "Desktop app"
#   3. Download the resulting JSON file, save it anywhere, and set its
#      full path as an environment variable: GOOGLE_CLIENT_SECRET_PATH
#
# First run opens your browser for a one-time login/consent. After that,
# the token is cached to disk and silently refreshed.
#
# Scope choice: gmail.modify + gmail.send covers everything our tools
# actually do (read, send, move to trash) without requesting the
# broader permanent-delete capability our delete_gmail_message doesn't
# use anyway (it moves to trash, not permanent deletion).
# ---------------------------------------------------------------------------

GMAIL_SCOPES = [
    "https://www.googleapis.com/auth/gmail.modify",
    "https://www.googleapis.com/auth/gmail.send",
]
GMAIL_BASE = "https://gmail.googleapis.com/gmail/v1"

_gmail_creds = None


def _get_gmail_token_path() -> str:
    return os.path.join(_get_data_dir(), "gmail_token.json")


def _get_gmail_token() -> str:
    """Same shape as _get_graph_token for Outlook: load a cached session
    if there is one, refresh it silently if it's expired, and only open a
    browser for a fresh login if there's no usable cached session at all."""
    global _gmail_creds

    token_path = _get_gmail_token_path()

    if _gmail_creds is None and os.path.exists(token_path):
        _gmail_creds = Credentials.from_authorized_user_file(token_path, GMAIL_SCOPES)

    if _gmail_creds and _gmail_creds.expired and _gmail_creds.refresh_token:
        _gmail_creds.refresh(GoogleAuthRequest())

    if not _gmail_creds or not _gmail_creds.valid:
        client_secret_path = os.environ.get("GOOGLE_CLIENT_SECRET_PATH")
        if not client_secret_path:
            raise RuntimeError("GOOGLE_CLIENT_SECRET_PATH environment variable is not set.")
        flow = InstalledAppFlow.from_client_secrets_file(client_secret_path, GMAIL_SCOPES)
        _gmail_creds = flow.run_local_server(port=0)

    with open(token_path, "w") as f:
        f.write(_gmail_creds.to_json())

    return _gmail_creds.token


def _gmail_request(method: str, path: str, params: dict = None, json_body: dict = None) -> requests.Response:
    token = _get_gmail_token()
    headers = {"Authorization": f"Bearer {token}"}
    return requests.request(method, f"{GMAIL_BASE}{path}", headers=headers, params=params, json=json_body, timeout=15)


def _b64url_decode(data: str) -> str:
    padded = data + "=" * (-len(data) % 4)
    return base64.urlsafe_b64decode(padded).decode("utf-8", errors="replace")


def _extract_gmail_body(payload: dict) -> str:
    """
    Gmail message bodies aren't a simple text field like Outlook's - they're
    a MIME structure that can be a single part or nested multipart/alternative
    parts (e.g. text/plain alongside text/html). This recursively hunts for
    a text/plain part and base64url-decodes it, since that's the most
    reliably readable form for ORACLE to work with.
    """
    if payload.get("mimeType") == "text/plain" and payload.get("body", {}).get("data"):
        return _b64url_decode(payload["body"]["data"])

    for part in payload.get("parts") or []:
        found = _extract_gmail_body(part)
        if found:
            return found

    # No text/plain part found anywhere - fall back to whatever's at the
    # top level, if anything (better than returning nothing).
    if payload.get("body", {}).get("data"):
        return _b64url_decode(payload["body"]["data"])

    return ""


def list_gmail_messages(count: int = 10) -> str:
    """Lists the most recent Gmail messages. Slower than Outlook's
    equivalent by design - Gmail's list endpoint only returns IDs, so
    this makes one follow-up request per message to get subject/sender."""
    try:
        resp = _gmail_request("GET", "/users/me/messages", params={"maxResults": count})
        resp.raise_for_status()
    except Exception as e:
        return f"Error listing Gmail messages: {e}"

    ids = [m["id"] for m in resp.json().get("messages", [])]
    if not ids:
        return "No messages found."

    lines = []
    for msg_id in ids:
        try:
            detail_resp = _gmail_request(
                "GET", f"/users/me/messages/{msg_id}",
                params={"format": "metadata", "metadataHeaders": ["Subject", "From", "Date"]},
            )
            detail_resp.raise_for_status()
            detail = detail_resp.json()
        except Exception:
            continue

        headers = {h["name"]: h["value"] for h in detail.get("payload", {}).get("headers", [])}
        lines.append(
            f"ID: {msg_id}\nFrom: {headers.get('From', 'unknown')}\n"
            f"Subject: {headers.get('Subject', '(no subject)')}\nDate: {headers.get('Date', '')}\n"
            f"Preview: {detail.get('snippet', '')[:200]}"
        )

    return "\n\n".join(lines) if lines else "No messages could be retrieved."


def read_gmail_message(message_id: str) -> str:
    """Reads the full content of a specific Gmail message. Get the
    message's ID from list_gmail_messages first."""
    try:
        resp = _gmail_request("GET", f"/users/me/messages/{message_id}", params={"format": "full"})
        resp.raise_for_status()
    except Exception as e:
        return f"Error reading Gmail message: {e}"

    msg = resp.json()
    headers = {h["name"]: h["value"] for h in msg.get("payload", {}).get("headers", [])}
    body = _extract_gmail_body(msg.get("payload", {}))
    return f"From: {headers.get('From', 'unknown')}\nSubject: {headers.get('Subject', '(no subject)')}\n\n{body}"


def send_gmail_message(to: str, subject: str, body: str) -> str:
    """Sends an email from the user's Gmail account. Sends immediately -
    there is no draft/confirmation step."""
    message = MIMEText(body)
    message["to"] = to
    message["subject"] = subject
    raw = base64.urlsafe_b64encode(message.as_bytes()).decode("utf-8")

    try:
        resp = _gmail_request("POST", "/users/me/messages/send", json_body={"raw": raw})
        resp.raise_for_status()
        return f"Email sent to {to}."
    except Exception as e:
        return f"Error sending email: {e}"


def delete_gmail_message(message_id: str) -> str:
    """Moves a Gmail message to Trash by ID (not permanently deleted -
    Gmail keeps trashed messages for 30 days before erasing them)."""
    try:
        resp = _gmail_request("POST", f"/users/me/messages/{message_id}/trash")
        resp.raise_for_status()
        return f"Moved message {message_id} to Trash."
    except Exception as e:
        return f"Error deleting Gmail message: {e}"


# This is the "menu" we hand to the model, describing each tool so it knows
# when and how to call it. This schema format is the same shape used by
# OpenAI, Anthropic, and Ollama - it's becoming a de facto standard.
TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "get_current_time",
            "description": "Get the current date and time.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_files",
            "description": "List files in a directory on the local machine.",
            "parameters": {
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "Path to the directory to list. Defaults to current directory.",
                    }
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_file",
            "description": (
                "Open a specific FILE (e.g. PDF, DOCX, image, text file) using the "
                "operating system's default application for that file type - "
                "the same as double-clicking it in a file browser. For launching "
                "an APPLICATION by name (e.g. Notepad, Chrome, Spotify) instead "
                "of opening a file, use launch_app instead - it's much faster."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Full or relative path to the file to open.",
                    }
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "launch_app",
            "description": (
                "Launch an application by name (e.g. 'notepad', 'chrome', 'spotify', "
                "'calculator'). Use this for opening PROGRAMS - it's a single fast "
                "call. Do NOT use list_files or open_file to hunt for an application's "
                "executable file first; launch_app resolves common app names directly."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "app_name": {
                        "type": "string",
                        "description": "Name of the application to launch, e.g. 'notepad' or 'chrome'.",
                    }
                },
                "required": ["app_name"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": (
                "Read and extract the text content of a file so you can summarize it "
                "or answer questions about it. Supports .txt, .md, .csv, .json, .log, "
                ".pdf, and .docx."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "Full or relative path to the file to read.",
                    }
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "move_file",
            "description": "Move or rename a file on the local machine.",
            "parameters": {
                "type": "object",
                "properties": {
                    "source": {"type": "string", "description": "Current path of the file."},
                    "destination": {"type": "string", "description": "New path, name, or destination directory."},
                },
                "required": ["source", "destination"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delete_file",
            "description": "Permanently delete a file from the local machine. This cannot be undone.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Full or relative path to the file to delete."}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_file",
            "description": "Create a new text file with the given content on the local machine.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Full or relative path for the new file."},
                    "content": {"type": "string", "description": "Text content to write into the file. Defaults to empty."},
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_system_info",
            "description": "Get current system resource usage: CPU, RAM, disk/storage, and GPU/VRAM if available.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ask_coding_agent",
            "description": (
                "Consult a coding-specialist model for non-trivial programming help - "
                "writing code, debugging, explaining code, architecture/design questions. "
                "Use this instead of answering coding questions yourself when the request "
                "is more than trivial (a one-line syntax question is fine to answer "
                "directly; writing/debugging/explaining real code should go through this)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {
                        "type": "string",
                        "description": "The coding question or task, including any necessary context (language, relevant code, error messages) since this doesn't see the rest of the conversation.",
                    }
                },
                "required": ["prompt"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": (
                "Search the web for current information and get results back as text "
                "for you to read and summarize in the conversation - the user does NOT "
                "see a browser. Use this when the user is asking a question you need "
                "an answer to. For when the user wants to actually SEE search results "
                "themselves in their browser, use open_web_search instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "The search query."}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_web_search",
            "description": (
                "Opens the user's default web browser to a real search results page "
                "they can see and interact with - use this when the user asks you to "
                "search for something and clearly wants to look at the results "
                "themselves (e.g. 'search for X', 'look up Y for me'), as opposed to "
                "asking you a question you should just answer directly (use web_search "
                "for that instead). If the user instead names a specific website to "
                "visit (e.g. 'open youtube.com', 'go to github'), use open_url instead."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "The search query."}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "open_url",
            "description": (
                "Opens a specific website in the user's default browser - use this "
                "when the user names an actual site or address to go to (e.g. 'open "
                "youtube.com', 'go to github.com/anthropics'), as opposed to searching "
                "for something (use web_search or open_web_search for that instead)."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "The web address to open, e.g. 'youtube.com' or 'https://github.com'.",
                    }
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "send_notification",
            "description": (
                "Show a real Windows toast/system notification - separate from "
                "the chat window, visible even if the user isn't looking at "
                "ORACLE right now. Use this for alerts the user should notice "
                "even if they're not in the app."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {"type": "string", "description": "Short notification title."},
                    "message": {"type": "string", "description": "Notification body text. Optional."},
                },
                "required": ["title"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_recent_emails",
            "description": "List the most recent emails in the user's Outlook inbox.",
            "parameters": {
                "type": "object",
                "properties": {
                    "count": {"type": "integer", "description": "How many recent emails to list. Defaults to 10."}
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_email",
            "description": "Read the full content of a specific Outlook email. Get the email's ID from list_recent_emails first.",
            "parameters": {
                "type": "object",
                "properties": {
                    "email_id": {"type": "string", "description": "The ID of the email to read."}
                },
                "required": ["email_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "send_email",
            "description": "Send an email from the user's Outlook account. Sends immediately.",
            "parameters": {
                "type": "object",
                "properties": {
                    "to": {"type": "string", "description": "Recipient email address."},
                    "subject": {"type": "string", "description": "Email subject line."},
                    "body": {"type": "string", "description": "Email body text."},
                },
                "required": ["to", "subject", "body"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delete_email",
            "description": "Delete an Outlook email by ID. Moves it to Deleted Items.",
            "parameters": {
                "type": "object",
                "properties": {
                    "email_id": {"type": "string", "description": "The ID of the email to delete."}
                },
                "required": ["email_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_upcoming_events",
            "description": "List upcoming events on the user's Outlook calendar.",
            "parameters": {
                "type": "object",
                "properties": {
                    "days": {"type": "integer", "description": "How many days ahead to look. Defaults to 7."}
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_calendar_event",
            "description": "Create a new event on the user's Outlook calendar.",
            "parameters": {
                "type": "object",
                "properties": {
                    "subject": {"type": "string", "description": "Event title."},
                    "start_iso": {"type": "string", "description": "Start time in ISO 8601 format, e.g. '2026-09-05T14:00:00'."},
                    "end_iso": {"type": "string", "description": "End time in ISO 8601 format."},
                    "attendees": {"type": "string", "description": "Optional comma-separated list of attendee email addresses."},
                },
                "required": ["subject", "start_iso", "end_iso"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delete_calendar_event",
            "description": "Delete/cancel an Outlook calendar event by ID. Cannot be undone.",
            "parameters": {
                "type": "object",
                "properties": {
                    "event_id": {"type": "string", "description": "The ID of the event to delete."}
                },
                "required": ["event_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_gmail_messages",
            "description": "List the most recent messages in the user's Gmail inbox (separate from Outlook - use this specifically for Gmail).",
            "parameters": {
                "type": "object",
                "properties": {
                    "count": {"type": "integer", "description": "How many recent messages to list. Defaults to 10."}
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_gmail_message",
            "description": "Read the full content of a specific Gmail message. Get the message's ID from list_gmail_messages first.",
            "parameters": {
                "type": "object",
                "properties": {
                    "message_id": {"type": "string", "description": "The ID of the message to read."}
                },
                "required": ["message_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "send_gmail_message",
            "description": "Send an email from the user's Gmail account. Sends immediately.",
            "parameters": {
                "type": "object",
                "properties": {
                    "to": {"type": "string", "description": "Recipient email address."},
                    "subject": {"type": "string", "description": "Email subject line."},
                    "body": {"type": "string", "description": "Email body text."},
                },
                "required": ["to", "subject", "body"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delete_gmail_message",
            "description": "Move a Gmail message to Trash by ID (recoverable for 30 days, not permanent).",
            "parameters": {
                "type": "object",
                "properties": {
                    "message_id": {"type": "string", "description": "The ID of the message to delete."}
                },
                "required": ["message_id"],
            },
        },
    },
]

# Map tool names to actual Python functions so we can execute them by name.
AVAILABLE_FUNCTIONS = {
    "get_current_time": get_current_time,
    "list_files": list_files,
    "open_file": open_file,
    "launch_app": launch_app,
    "read_file": read_file,
    "move_file": move_file,
    "delete_file": delete_file,
    "create_file": create_file,
    "get_system_info": get_system_info,
    "ask_coding_agent": ask_coding_agent,
    "web_search": web_search,
    "open_web_search": open_web_search,
    "open_url": open_url,
    "send_notification": send_notification,
    "list_recent_emails": list_recent_emails,
    "read_email": read_email,
    "send_email": send_email,
    "delete_email": delete_email,
    "list_upcoming_events": list_upcoming_events,
    "create_calendar_event": create_calendar_event,
    "delete_calendar_event": delete_calendar_event,
    "list_gmail_messages": list_gmail_messages,
    "read_gmail_message": read_gmail_message,
    "send_gmail_message": send_gmail_message,
    "delete_gmail_message": delete_gmail_message,
}


def trim_history(history: list) -> list:
    """
    Keeps the system prompt plus only the most recent MAX_HISTORY_MESSAGES
    messages. Everything is still saved in full to SQLite - this only
    controls what gets sent to the API each turn, since token usage (and
    Groq's rate limits) scale with how much history you resend every call.
    """
    system_msgs = [m for m in history if m["role"] == "system"]
    other_msgs = [m for m in history if m["role"] != "system"]
    trimmed = other_msgs[-MAX_HISTORY_MESSAGES:]

    # A "tool" message only makes sense immediately after the assistant
    # message that requested it. If the cut landed between them, the API
    # will reject the orphaned tool message - so drop leading tool messages
    # until we start on a clean boundary.
    while trimmed and trimmed[0]["role"] == "tool":
        trimmed.pop(0)

    return system_msgs + trimmed


def run_conversation(user_input: str, history: list, conversation_id: int = None) -> str:
    """
    Sends the user's message + history to the model. The model may need several
    rounds of tool calls - each result can prompt the next call - so we keep
    going until it returns a plain text answer. conversation_id tags every
    saved message to the right thread so it shows up correctly if this
    conversation is reopened later from the sidebar.
    """
    history[:] = trim_history(history)
    user_msg = {"role": "user", "content": user_input}
    history.append(user_msg)
    save_message(user_msg, conversation_id)

    for _ in range(MAX_TOOL_ROUNDS):
        response = client.chat.completions.create(
            model=MODEL,
            messages=history,
            tools=TOOLS,
        )

        message = response.choices[0].message

        # No tool needed - just a normal reply.
        if not message.tool_calls:
            final_msg = {"role": "assistant", "content": message.content}
            history.append(final_msg)
            save_message(final_msg, conversation_id)
            return message.content

        # Record the assistant's tool-call request in history. We build this
        # dict manually rather than using message.model_dump() - the full
        # dump includes extra fields (like "annotations") that the API is
        # happy to SEND but refuses to RECEIVE back. Only include what the
        # spec actually requires.
        tool_call_msg = {
            "role": "assistant",
            "content": message.content,
            "tool_calls": [
                {
                    "id": tc.id,
                    "type": "function",
                    "function": {
                        "name": tc.function.name,
                        "arguments": tc.function.arguments,
                    },
                }
                for tc in message.tool_calls
            ],
        }
        history.append(tool_call_msg)
        save_message(tool_call_msg, conversation_id)

        for tool_call in message.tool_calls:
            fn_name = tool_call.function.name
            # Arguments come back as a JSON string, not a dict - must parse.
            fn_args = json.loads(tool_call.function.arguments)

            fn = AVAILABLE_FUNCTIONS.get(fn_name)
            result = fn(**fn_args) if fn else f"Unknown tool: {fn_name}"

            # Feed the tool's result back to the model as a "tool" message.
            # tool_call_id links this result to the specific call above -
            # required when a model requests multiple tool calls at once.
            tool_result_msg = {
                "role": "tool",
                "tool_call_id": tool_call.id,
                "content": str(result),
            }
            history.append(tool_result_msg)
            save_message(tool_result_msg, conversation_id)

    give_up = f"I kept calling tools without reaching an answer ({MAX_TOOL_ROUNDS} rounds). Try asking a narrower question."
    give_up_msg = {"role": "assistant", "content": give_up}
    history.append(give_up_msg)
    save_message(give_up_msg, conversation_id)
    return give_up


def generate_wake_greeting() -> str:
    """
    Kept in case wake-word or a similar hands-free activation trigger
    comes back later - not currently called anywhere, since clicking the
    HUD ring goes straight to listening without a greeting step.
    """
    prompt = (
        f"The current time is {get_current_time()}. The user (whom you address as "
        f"'{USER_NAME}') just activated you. Greet them by name "
        "in one short, natural sentence, in character as established in your system "
        "prompt - calm, dry-witted, quietly loyal. Vary your phrasing meaningfully "
        "each time rather than repeating a fixed template - not a generic "
        "'How can I help you today?'"
    )
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
    )
    return response.choices[0].message.content


if __name__ == "__main__":
    # cp1252, the default Windows console encoding, can't represent characters
    # that routinely appear in replies (curly quotes, U+202F).
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    init_db()

    print("ORACLE core - terminal test mode. Type 'quit' to exit.\n")
    print("(For the full ORACLE experience, run ui.py instead.)\n")

    conversation_history = [{"role": "system", "content": SYSTEM_PROMPT}]
    prior_messages = load_recent_history()
    if prior_messages:
        conversation_history.extend(prior_messages)

    while True:
        user_text = input("You: ")
        if user_text.lower() in ("quit", "exit"):
            break
        reply = run_conversation(user_text, conversation_history)
        print(f"ORACLE: {reply}\n")